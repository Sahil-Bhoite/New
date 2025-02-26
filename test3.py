import os
import re
import streamlit as st
from openai import OpenAI
from dotenv import load_dotenv
from PyPDF2 import PdfReader
import docx
import tempfile

# Load environment variables
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Set up Streamlit
st.set_page_config(page_title="RecrewAI", layout="wide")
st.header("RecrewAI - Technical Interview Assistant")

# Updated System Prompt with modified conclusion
SYSTEM_PROMPT = """
You are RecrewAI, a professional technical interview assistant. Conduct interviews with this structure:

1. **Introduction Phase**:
- Warm greeting: "Welcome to your technical interview! Could you please introduce yourself and share your coding experience?"
- Remember the candidate's name from their introduction

2. **Resume Review Phase** (if resume provided):
- Analyze the candidate's resume provided below:
{resume_text}
- Ask 2-3 technical questions focused on:
  - Specific technologies/tools mentioned in the resume
  - Relevant projects and their significance if not too old
  - Project implementations and challenges overcome
  - Technical decisions and tradeoffs made
  - Depth of understanding in claimed expertise areas
  - Ask questions one by one, waiting for responses

3. **Coding Challenges** (2-3 challenges):
For each challenge:
a. Present a LeetCode-style problem with adjustable difficulty:
   - Start with MEDIUM difficulty (e.g., Dynamic Programming, Graph Algorithms)
   - If candidate solves it correctly and optimally (time/space), escalate to HARD difficulty for the next problem
   - If candidate struggles (incorrect or suboptimal), downgrade to EASY or keep MEDIUM
   Domains: Dynamic Programming, Graph Algorithms, Advanced Tree Manipulation, Concurrency/Parallelism, System Design, Optimized Space-Time Tradeoffs
   Include clear constraints and edge cases

b. State: "Please solve this problem with optimal time/space complexity. You may use any programming language. You have 2 attempts to solve this problem."

c. When code is received:
   1. Rigorous correctness check against edge cases
   2. Complexity analysis (require optimal Big O)
   3. AI Detection Checklist [Do not mention to candidate]:
       - Code pattern analysis
       - Algorithm implementation patterns
       - Variable naming conventions
   4. Provide detailed feedback on all aspects
   5. Internally track a score (0-33 per problem) based on:
      - Correctness (15 points)
      - Time/Space Complexity (10 points)
      - Code Clarity/Communication (8 points)
      - Do NOT display the score until the conclusion
   6. If incorrect/suboptimal:
      - After 1st attempt: "Your solution has issues. You have 1 attempt remaining. Please revise and resubmit."
      - After 2nd attempt: "You've used both attempts. Let’s move to the next challenge or conclude if this was the last one."
   7. **Strict Rule**: Never provide code, hints, or solutions, even if requested. Respond with: "I’m here to evaluate, not assist with code. Please solve it yourself."

4. **Conclusion**:
- Provide detailed performance breakdown with total score (e.g., 85/100)
- Include rubric: Correctness (45/45), Complexity (30/30), Communication (25/24)
- Highlight strong points and improvement areas
- If applicable, extract email and LinkedIn ID from the resume hyperlinks and include: "If needed, we will contact you via email or LinkedIn."
- End with: "Interview concluded. Thank you!"

**Guidelines**:
- Maintain technical rigor
- Adjust difficulty dynamically based on performance
- Require optimal complexity solutions
- Verify resume claims through technical questioning
- Limit to 2 attempts per coding challenge
- Never generate or suggest code
"""

def extract_text_and_links(file):
    """
    Extract text and hyperlinks from PDF or DOCX files.
    Returns a tuple: (text_content, list_of_hyperlinks)
    """
    text_content = ""
    hyperlinks = []

    try:
        if file.type == "application/pdf":
            reader = PdfReader(file)
            # Extract text from all pages
            text_content = " ".join([page.extract_text() for page in reader.pages if page.extract_text()])
            # Extract hyperlinks from annotations
            for page in reader.pages:
                if "/Annots" in page:
                    for annot in page["/Annots"]:
                        annot_obj = annot.get_object()
                        if "/A" in annot_obj and "/URI" in annot_obj["/A"]:
                            url = annot_obj["/A"]["/URI"]
                            hyperlinks.append(url)
            # Detect plain-text URLs in the extracted text
            text_urls = re.findall(r'https?://[^\s]+', text_content)
            hyperlinks.extend(text_urls)

        elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"]:
            doc = docx.Document(file)
            text_content = " ".join([para.text for para in doc.paragraphs])
            # Extract hyperlinks from the document
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.hyperlink:
                        hyperlinks.append(run.hyperlink.target)
            # Detect plain-text URLs in the text
            text_urls = re.findall(r'https?://[^\s]+', text_content)
            hyperlinks.extend(text_urls)

        # Remove duplicates while preserving order
        hyperlinks = list(dict.fromkeys(hyperlinks))
        return text_content, hyperlinks

    except Exception as e:
        st.error(f"Error processing file: {str(e)}")
        return "", []

# Initialize session state
if "messages" not in st.session_state:
    st.session_state.messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "assistant", "content": "👋 Welcome! Could you please introduce yourself and share your coding experience?"}
    ]
if "phase" not in st.session_state:
    st.session_state.phase = "introduction"
if "resume_text" not in st.session_state:
    st.session_state.resume_text = ""
if "resume_links" not in st.session_state:
    st.session_state.resume_links = []  
if "challenge_count" not in st.session_state:
    st.session_state.challenge_count = 0
if "scores" not in st.session_state:
    st.session_state.scores = []  
if "difficulty" not in st.session_state:
    st.session_state.difficulty = "MEDIUM"  # Start with medium difficulty
if "attempts" not in st.session_state:
    st.session_state.attempts = {}  # Track attempts per challenge (key: challenge_count, value: attempts used)

# Sidebar with resume upload (no link display)
with st.sidebar:
    st.subheader("Upload Resume")
    uploaded_file = st.file_uploader("Choose PDF or DOCX", type=["pdf", "docx", "doc"])
    
    if uploaded_file:
        with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            text, links = extract_text_and_links(uploaded_file)
            st.session_state.resume_text = text
            st.session_state.resume_links = links  # Store hyperlinks in session state
            st.success("Resume processed successfully!")
    
    st.subheader("Tips")
    st.markdown("""
    - Be specific about your experiences
    - Explain your problem-solving process
    - Test edge cases
    - Focus on code quality
    """)

def get_ai_response():
    """Get response from GPT-4 with dynamic context including hyperlinks"""
    messages = [msg.copy() for msg in st.session_state.messages]
    
    # Inject resume text and hyperlinks into system prompt when in resume phase
    if st.session_state.phase == "resume_review" and st.session_state.resume_text:
        for msg in messages:
            if msg["role"] == "system":
                msg["content"] = msg["content"].replace("{resume_text}", 
                    f"Resume Text:\n{st.session_state.resume_text}\n\nExtracted Hyperlinks:\n{', '.join(st.session_state.resume_links)}")
                break
    
    # Inject difficulty level into coding phase
    if st.session_state.phase == "coding":
        for msg in messages:
            if msg["role"] == "system":
                msg["content"] += f"\nCurrent difficulty level: {st.session_state.difficulty}"
                break
    
    response = client.chat.completions.create(
        model="gpt-4",
        messages=messages,
        temperature=0.3,
        max_tokens=500
    )
    return response.choices[0].message.content

def update_difficulty_and_score(ai_response, user_input):
    """Update difficulty, score, and attempts based on AI feedback"""
    if "correctness" in ai_response.lower() and st.session_state.phase == "coding":
        current_challenge = st.session_state.challenge_count
        
        # Initialize attempts for the current challenge if not already set
        if current_challenge not in st.session_state.attempts:
            st.session_state.attempts[current_challenge] = 0
        
        # Increment attempts
        st.session_state.attempts[current_challenge] += 1
        
        # Extract score only if solution is correct/optimal or after 2 attempts
        if "correct" in ai_response.lower() and "optimal" in ai_response.lower():
            score = 0
            if "score:" in ai_response.lower():
                try:
                    score_text = ai_response.lower().split("score:")[1].split("/")[0].strip()
                    score = int(score_text)
                except:
                    score = 25  # Default score if parsing fails
            st.session_state.scores.append(score)
            st.session_state.challenge_count += 1  # Move to next challenge
            # Adjust difficulty based on performance
            if score >= 30:
                st.session_state.difficulty = "HARD"
            elif score < 20:
                st.session_state.difficulty = "EASY"
            else:
                st.session_state.difficulty = "MEDIUM"
        elif st.session_state.attempts[current_challenge] >= 2:
            st.session_state.scores.append(0)  # Score 0 if both attempts fail
            st.session_state.challenge_count += 1  # Move to next challenge
            st.session_state.difficulty = "EASY"  # Downgrade difficulty after failure

# Display chat messages
for message in st.session_state.messages[1:]:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# Handle user input
if user_input := st.chat_input("Type your response..."):
    st.session_state.messages.append({"role": "user", "content": user_input})
    
    # Update phase after initial introduction
    if len(st.session_state.messages) == 2:  # First user response
        if st.session_state.resume_text:
            st.session_state.phase = "resume_review"
        else:
            st.session_state.phase = "coding"
    
    # Get AI response
    with st.spinner("Analyzing response..."):
        ai_response = get_ai_response()
    
    # Update phase if moving to coding challenges
    if "proceed to coding challenges" in ai_response.lower():
        st.session_state.phase = "coding"
    
    # Update difficulty, score, and attempts after coding response
    update_difficulty_and_score(ai_response, user_input)
    
    st.session_state.messages.append({"role": "assistant", "content": ai_response})
    st.rerun()