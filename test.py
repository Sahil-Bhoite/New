import os
import re
import streamlit as st
from openai import OpenAI
from dotenv import load_dotenv
from PyPDF2 import PdfReader
import docx
import tempfile

# Load environment variables
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    st.error("API key not found. Please set OPENAI_API_KEY in your .env file.")
    st.stop()
client = OpenAI(api_key=api_key)

# Set up Streamlit
st.set_page_config(page_title="RecrewAI", layout="wide")
st.header("RecrewAI - Technical Interview Assistant")

# System Prompt with Emojis and Natural Tone
SYSTEM_PROMPT = """
You are RecrewAI, a chill and supportive technical interview buddy. Keep it fun, natural, and fluent like a real chat. Here’s the flow:

1. **Introduction Phase**:
- Kick off with: "Hey! Welcome to your tech chat 😊 Tell me about yourself and what coding stuff you’ve been up to!"
- Grab their name and sprinkle it in to feel personal.

2. **Resume Review Phase** (if resume provided):
- Peek at: {resume_text}
- Toss out 2-3 easy questions like:
  - "What’s a tool you’ve messed with that you liked?"
  - "Got a project you’re stoked about? What was the hard part?"
  - "Any big choices you made on a project?"
- Go one-by-one, vibe with their answers.
- Use ✅ for "yep, cool!" (e.g., "✅ Love that!") and ❌ for "hmm, not quite" (e.g., "❌ Might need a tweak—why’d you pick that?").
- After 2-3, ask: "Ready to jump into some coding?"

3. **Coding Prep Phase**:
- If they say "yes" to "Ready to jump into some coding?", say: "Cool, [name]! Here’s your first coding challenge: [present a MEDIUM difficulty problem with constraints and edge cases]."
- If "no," say: "No rush—what’s up? How can I get you pumped for coding?" and nudge them to "yes."

4. **Coding Challenges** (2-3 challenges):
For each one:
a. After presenting the problem, ask: "Any quick questions to sort this out?"
   - Answer like a pal, e.g., "✅ Good one! Here’s the deal..."

b. Say: "Sweet—how’d you solve this?"
   - Chat it out. Use ✅ for "nice!" (e.g., "✅ That’s sharp!") and ❌ for "oops" (e.g., "❌ Could miss this—what about...?").
   - When it’s solid, say: "✅ Awesome—code it up!"

c. When they give code:
   - Check it hard: correctness, edges, complexity.
   - Throw 2-3 chill follow-ups like:
     - "Can we juice this up more?"
     - "Time and space—what’s the scoop?"
     - "Any weird cases this skips?"
   - Use ✅/❌ (e.g., "✅ Nailed it!" or "❌ Hmm, misses negatives—thoughts?").
   - Let them tweak it till it shines, then move on.
   - Score it quietly (0-33 per problem):
     - Correctness: 15
     - Complexity: 10
     - Clarity/Chat: 8
   - Adjust difficulty: HARD if they ace it, EASY if they stumble.

5. **Conclusion**:
- Wrap with: "You rocked this! Here’s the rundown: [score]/100"
- Split it: Correctness (X/45), Complexity (Y/30), Chat (Z/24)
- Say what’s great and what to polish.
- If resume’s there, grab email/LinkedIn and say: "Might ping you later via email or LinkedIn!"
- End: "That’s a wrap—thanks for chilling with me!"

**Guidelines**:
- Act like a friend who knows their stuff
- Keep it upbeat with ✅ and ❌ for quick feedback
- No code or hints—just fun questions to dig deeper
"""

# Function to extract text and links from resume
def extract_text_and_links(file):
    text_content = ""
    hyperlinks = []
    try:
        if file.type == "application/pdf":
            reader = PdfReader(file)
            text_content = " ".join([page.extract_text() for page in reader.pages if page.extract_text()])
            for page in reader.pages:
                if "/Annots" in page:
                    for annot in page["/Annots"]:
                        annot_obj = annot.get_object()
                        if "/A" in annot_obj and "/URI" in annot_obj["/A"]:
                            hyperlinks.append(annot_obj["/A"]["/URI"])
            text_urls = re.findall(r'https?://[^\s]+', text_content)
            hyperlinks.extend(text_urls)
        elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            doc = docx.Document(file)
            text_content = " ".join([para.text for para in doc.paragraphs])
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.hyperlink:
                        hyperlinks.append(run.hyperlink.target)
            text_urls = re.findall(r'https?://[^\s]+', text_content)
            hyperlinks.extend(text_urls)
        hyperlinks = list(dict.fromkeys(hyperlinks))
        return text_content, hyperlinks
    except Exception as e:
        st.error(f"Error processing file: {str(e)}")
        return "", []

# Function to soften tone
def soften_tone(response):
    if "correctness" in response.lower():
        response = f"Love the effort! {response}"
    elif "what" in response.lower() and "?" in response:
        response = f"Fun question! {response}"
    elif "not quite" in response.lower() or "miss" in response.lower():
        response = f"No worries—let’s tweak it! {response}"
    return response.strip()

# Initialize session state
if "messages" not in st.session_state:
    st.session_state.messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "assistant", "content": "Hey! Welcome to your tech chat I am recrewAI. Tell me about yourself and what coding stuff you’ve been up to!"}
    ]
if "phase" not in st.session_state:
    st.session_state.phase = "introduction"
if "candidate_name" not in st.session_state:
    st.session_state.candidate_name = ""
if "resume_text" not in st.session_state:
    st.session_state.resume_text = ""
if "resume_links" not in st.session_state:
    st.session_state.resume_links = []
if "challenge_count" not in st.session_state:
    st.session_state.challenge_count = 0
if "scores" not in st.session_state:
    st.session_state.scores = []
if "difficulty" not in st.session_state:
    st.session_state.difficulty = "MEDIUM"
if "resume_questions_asked" not in st.session_state:
    st.session_state.resume_questions_asked = 0
if "coding_ready" not in st.session_state:
    st.session_state.coding_ready = False
if "awaiting_clarification" not in st.session_state:
    st.session_state.awaiting_clarification = False
if "awaiting_approach" not in st.session_state:
    st.session_state.awaiting_approach = False
if "awaiting_code" not in st.session_state:
    st.session_state.awaiting_code = False
if "problem_presented" not in st.session_state:
    st.session_state.problem_presented = False

# Sidebar with resume upload
with st.sidebar:
    st.subheader("Upload Resume")
    uploaded_file = st.file_uploader("Choose PDF or DOCX", type=["pdf", "docx", "doc"])
    if uploaded_file:
        with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            text, links = extract_text_and_links(uploaded_file)
            st.session_state.resume_text = text
            st.session_state.resume_links = links
            st.success("Resume’s in—sweet!")
    
    st.subheader("Tips")
    st.markdown("""
    - Brag a little about your wins
    - Walk me through your thinking
    - Watch those tricky edge cases
    - Keep it easy to follow
    """)

def get_ai_response():
    """Get response from GPT-4 with dynamic context"""
    messages = [msg.copy() for msg in st.session_state.messages]
    if st.session_state.phase == "resume_review" and st.session_state.resume_text:
        for msg in messages:
            if msg["role"] == "system":
                msg["content"] = msg["content"].replace("{resume_text}", 
                    f"Resume Text:\n{st.session_state.resume_text}\n\nExtracted Hyperlinks:\n{', '.join(st.session_state.resume_links)}")
                break
    if st.session_state.phase == "coding":
        for msg in messages:
            if msg["role"] == "system":
                msg["content"] += f"\nCurrent difficulty level: {st.session_state.difficulty}"
                break
    response = client.chat.completions.create(
        model="gpt-4",
        messages=messages,
        temperature=0.6,
        max_tokens=500
    )
    return soften_tone(response.choices[0].message.content)

def update_phase(ai_response):
    """Update phase based on AI response"""
    if "that’s a wrap" in ai_response.lower() or "interview concluded" in ai_response.lower():
        st.session_state.phase = "conclusion"

def handle_coding_transition(user_input, ai_response):
    """Handle transition to coding phase"""
    if "ready to jump into some coding" in ai_response.lower():
        if "yes" in user_input.lower() or "sure" in user_input.lower():
            st.session_state.coding_ready = True
            st.session_state.phase = "coding"
            st.session_state.challenge_count = 1
            st.session_state.problem_presented = True
            name = st.session_state.candidate_name or "you"
            return f"Cool, {name}! Here’s your first coding challenge: Given an integer array nums, find the contiguous subarray with the largest sum. Constraints: 1 <= nums.length <= 3 * 10^4, -10^5 <= nums[i] <= 10^5. Edge cases: all negatives, all positives."
        elif "no" in user_input.lower():
            st.session_state.coding_ready = False
    elif not st.session_state.coding_ready and "what’s up" in ai_response.lower():
        if "ready" in user_input.lower():
            st.session_state.coding_ready = True
            st.session_state.phase = "coding"
            st.session_state.challenge_count = 1
            st.session_state.problem_presented = True
            name = st.session_state.candidate_name or "you"
            return f"Cool, {name}! Here’s your first coding challenge: Given an integer array nums, find the contiguous subarray with the largest sum. Constraints: 1 <= nums.length <= 3 * 10^4, -10^5 <= nums[i] <= 10^5. Edge cases: all negatives, all positives."
    return ai_response

def ask_clarification_questions(ai_response):
    """Manage clarification questions for coding problems"""
    if (st.session_state.phase == "coding" and 
        st.session_state.challenge_count > 0 and 
        st.session_state.problem_presented and 
        not st.session_state.awaiting_clarification and 
        "coding challenge" in ai_response.lower()):
        st.session_state.awaiting_clarification = True
        return ai_response + "\n\nAny quick questions to sort this out?"
    return ai_response

def evaluate_approach(ai_response, user_input):
    """Evaluate candidate's approach and prompt for code"""
    if st.session_state.awaiting_clarification and "any quick questions to sort this out" in ai_response.lower():
        st.session_state.awaiting_clarification = False
        st.session_state.awaiting_approach = True
        return "Sweet—how’d you solve this?"
    elif st.session_state.awaiting_approach:
        if "awesome—code it up" in ai_response.lower() or "nice" in ai_response.lower():
            st.session_state.awaiting_approach = False
            st.session_state.awaiting_code = True
            st.session_state.problem_presented = False  # Reset for next problem
            return "✅ Awesome—code it up!"
    return ai_response

def update_difficulty_and_score(ai_response, user_input):
    """Update difficulty and score based on AI feedback"""
    if "correctness" in ai_response.lower() and st.session_state.phase == "coding" and st.session_state.awaiting_code:
        if "correct" in ai_response.lower() and "optimal" in ai_response.lower():
            score = 25
            if "score:" in ai_response.lower():
                try:
                    score = int(ai_response.lower().split("score:")[1].split("/")[0].strip())
                except:
                    pass
            st.session_state.scores.append(score)
            st.session_state.awaiting_code = False
            st.session_state.challenge_count += 1
            if score >= 30:
                st.session_state.difficulty = "HARD"
            elif score < 20:
                st.session_state.difficulty = "EASY"
            else:
                st.session_state.difficulty = "MEDIUM"

# Display chat messages
for message in st.session_state.messages[1:]:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# Handle user input
if user_input := st.chat_input("Type your response..."):
    st.session_state.messages.append({"role": "user", "content": user_input})
    
    # After introduction, grab name and move to next phase
    if st.session_state.phase == "introduction":
        st.session_state.candidate_name = user_input.split()[0]  # Simple name grab
        st.session_state.phase = "resume_review" if st.session_state.resume_text else "coding"
        st.session_state.challenge_count = 1 if not st.session_state.resume_text else 0
    
    # Get AI response
    with st.spinner("Chatting it up..."):
        ai_response = get_ai_response()
    
    # Resume review phase: ask warm-up questions
    if st.session_state.phase == "resume_review":
        st.session_state.resume_questions_asked += 1
        if st.session_state.resume_questions_asked >= 2:
            ai_response = "Ready to jump into some coding?"
    
    # Handle coding transition
    ai_response = handle_coding_transition(user_input, ai_response)
    
    # Coding phase: clarification and approach
    if st.session_state.phase == "coding":
        ai_response = ask_clarification_questions(ai_response)
        ai_response = evaluate_approach(ai_response, user_input)
        update_difficulty_and_score(ai_response, user_input)
    
    # Update phase for conclusion
    update_phase(ai_response)
    
    st.session_state.messages.append({"role": "assistant", "content": ai_response})
    st.rerun()